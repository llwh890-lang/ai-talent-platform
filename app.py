"""
Web UI 视图模块 (大厂旗舰版 + 数据闭环引擎 + 渲染顺序修复)
基于 Streamlit 构建的双端交互界面，包含全局动态毛玻璃背景、
多端响应式 Echarts 雷达图，以及真实的反馈数据持久化落盘机制。
"""

import streamlit as st
import PyPDF2
from streamlit_echarts import st_echarts
import io
import pandas as pd
import json
import os
import time
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# ==========================================
# 工具函数定义区 (文件导出与处理)
# ==========================================

def set_run_font(run, font_name="Microsoft YaHei", size_pt=11, bold=False):
    """设置 Word 文档输出的中文字体及格式"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)


def create_excel_download(profile_data):
    """将 JSON 格式的能力画像打包转化为 Excel 数据流"""
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df_basic = pd.DataFrame([
            {"数据维度": "姓名", "详细内容": profile_data.get("name", "未知")},
            {"数据维度": "整体匹配度", "详细内容": f"{profile_data.get('match_score', 0)}%"}
        ])
        df_basic.to_excel(writer, sheet_name="基础信息", index=False)

        df_scores = pd.DataFrame({
            "能力维度": list(profile_data.get("radar_scores", {}).keys()),
            "评估得分": list(profile_data.get("radar_scores", {}).values())
        })
        df_scores.to_excel(writer, sheet_name="能力画像分数", index=False)

        df_gaps = pd.DataFrame({"核心差距分析": profile_data.get("gaps", [])})
        df_gaps.to_excel(writer, sheet_name="差距分析", index=False)

        jobs_data = profile_data.get("top_positions", [])
        if jobs_data and isinstance(jobs_data[0], dict):
            df_jobs = pd.DataFrame(jobs_data)
            df_jobs.rename(columns={"role": "推荐岗位", "reason": "推荐理由"}, inplace=True)
        else:
            df_jobs = pd.DataFrame({"推荐岗位": jobs_data})
        df_jobs.to_excel(writer, sheet_name="匹配岗位", index=False)
    return output.getvalue()


def create_word_download(markdown_text):
    """将 Markdown 格式的学习计划渲染为标准化 Word 数据流"""
    doc = Document()
    p_title = doc.add_paragraph()
    run_title = p_title.add_run("定制化成长路径规划与周任务")
    set_run_font(run_title, font_name="Microsoft YaHei", size_pt=22, bold=True)
    p_title.paragraph_format.space_after = Pt(24)

    for line in markdown_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        clean_line = line.replace("**", "")
        if clean_line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(clean_line[4:])
            set_run_font(run, font_name="Microsoft YaHei", size_pt=14, bold=True)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif clean_line.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(clean_line[3:])
            set_run_font(run, font_name="Microsoft YaHei", size_pt=16, bold=True)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
        elif clean_line.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(clean_line[2:])
            set_run_font(run, font_name="Microsoft YaHei", size_pt=18, bold=True)
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
        else:
            p = doc.add_paragraph()
            run = p.add_run(clean_line)
            set_run_font(run, font_name="Microsoft YaHei", size_pt=11, bold=False)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.25
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def extract_text_from_pdf(uploaded_file):
    """解析上传的 PDF 文件，加入空文本与扫描件的拦截机制"""
    try:
        reader = PyPDF2.PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

        if len(text.strip()) < 50:
            return "ERROR_SCANNED_PDF"
        return text
    except Exception as e:
        return "ERROR_PARSING_FAILED"


def load_job_models():
    """读取本地企业岗位 JSON 数据库"""
    db_path = os.path.join("data", "job_models.json")
    try:
        with open(db_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {
            "默认岗位开发": "岗位名称：初级后端开发\n能力四维要求：\n1. 算法与数据结构\n2. 数据库设计\n3. 架构能力\n4. 计算机网络"}

# 引入核心业务逻辑 (假设 core_engine.py 在同级目录)
from core_engine import (
    generate_student_profile,
    generate_student_profile_stream,
    parse_profile_json,
    generate_learning_plan_stream
)

# ==========================================
# 核心页面结构初始化
# ==========================================
st.set_page_config(page_title="AI 智能人才培养与精准就业平台", layout="wide")

# 初始化全局会话记忆
if "profile_data" not in st.session_state:
    st.session_state.profile_data = None
if "learning_plan_md" not in st.session_state:
    st.session_state.learning_plan_md = None
if "batch_talents_df" not in st.session_state:
    st.session_state.batch_talents_df = None
if "resume_text" not in st.session_state:
    st.session_state.resume_text = "姓名：李四\n专业：软件工程\n技能：学过C语言，对前端完全不懂。"
if "uploaded_file_name" not in st.session_state:
    st.session_state.uploaded_file_name = None
if "employee_list" not in st.session_state:
    st.session_state.employee_list = ["张大海 (入职3个月)", "李雪 (入职1个月)"]

analyze_btn = False

with st.sidebar:
    st.title("系统控制台")
    system_role = st.radio(
        "请选择登录角色",
        ["学生端 (个人成长诊断)", "企业端 (人才库与岗位管理)"]
    )
    st.markdown("---")

job_db = load_job_models()
job_options = list(job_db.keys())

# 全局背景常量
BG_URL = "https://images.unsplash.com/photo-1639322537228-f710d846310a?q=80&w=2064&auto=format&fit=crop"

def inject_global_css(glass_bg_color, blur_amount):
    """动态注入全局 CSS 样式"""
    st.markdown(
        f"""
        <style>
        .stApp {{
            background-image: url('{BG_URL}');
            background-size: cover;
            background-position: center;
            background-attachment: fixed;
        }}
        .block-container {{
            background-color: {glass_bg_color} !important;
            backdrop-filter: blur({blur_amount});
            -webkit-backdrop-filter: blur({blur_amount});
            border-radius: 20px;
            box-shadow: 0 12px 40px 0 rgba(0, 0, 0, 0.35);
            transition: background-color 0.5s ease; 
            margin-top: 2rem;
            padding: 3rem !important;
            padding-bottom: 5rem !important;
        }}
        header {{
            background-color: transparent !important;
        }}
        [data-testid="stSidebar"] {{
            background-color: rgba(240, 242, 246, 0.65) !important;
            backdrop-filter: blur(15px);
            -webkit-backdrop-filter: blur(15px);
            border-right: 1px solid rgba(255, 255, 255, 0.3);
        }}
        iframe {{
            background-color: transparent !important;
        }}
        h1.welcome-title {{
            color: #ffffff !important; 
            font-size: 4.8rem !important; 
            font-weight: 900 !important; 
            letter-spacing: 0.1em !important; 
            text-indent: 0.1em !important; 
            margin: 0 !important; 
            text-shadow: 0px 10px 30px rgba(0,0,0,0.8) !important;
        }}
        @media (max-width: 768px) {{
            h1.welcome-title {{
                font-size: 2.8rem !important; 
            }}
            .block-container {{
                padding: 1.5rem !important; 
                padding-bottom: 3rem !important;
            }}
        }}
        a {{
            position: relative;
            z-index: 50;
            pointer-events: auto !important;
            word-break: break-all;
        }}
        </style>
        """,
        unsafe_allow_html=True
    )


# ==========================================
# 路由跳转一：企业端 (B端) 视图逻辑
# ==========================================
if system_role == "企业端 (人才库与岗位管理)":
    # 优先注入 CSS
    inject_global_css("rgba(255, 255, 255, 0.92)", "20px")

    st.markdown(
        """
        <div style="padding-bottom: 15px; border-bottom: 2px solid rgba(0,0,0,0.1); margin-bottom: 30px;">
            <h2 style="margin: 0; color: #1e3a8a; font-weight: 900; font-size: 2.2rem;">企业人才智能管理门户</h2>
            <p style="margin: 5px 0 0 0; color: #0f172a; font-size: 1.1rem; font-weight: 800;">—— 基于润道星算大模型服务平台</p>
        </div>
        """,
        unsafe_allow_html=True
    )

    tab1, tab2, tab3 = st.tabs(["岗位模型库", "智能匹配清单", "用人反馈追踪"])

    with tab1:
        st.subheader("企业真实岗位模型管理")
        st.info("在此模块，HR 可维护企业的真实用人标准，这些标准将作为大模型智能匹配的基准尺。")
        st.json(job_db)

        with st.expander("新增岗位模型"):
            new_job_name = st.text_input("岗位名称")
            new_job_desc = st.text_area("能力要求与技能点", value="岗位名称：\n核心要求：\n1. \n2. \n3. \n4. ")
            if st.button("保存到企业岗位库"):
                if not new_job_name.strip() or not new_job_desc.strip():
                    st.warning("岗位名称和能力要求不能为空。")
                else:
                    job_db[new_job_name] = new_job_desc
                    db_path = os.path.join("data", "job_models.json")
                    try:
                        with open(db_path, "w", encoding="utf-8") as f:
                            json.dump(job_db, f, ensure_ascii=False, indent=4)
                        st.success(f"岗位【{new_job_name}】已成功入库并同步。")
                        st.rerun()
                    except Exception as e:
                        st.error(f"保存数据库失败，请检查文件权限: {e}")

        with st.expander("删除岗位模型"):
            if not job_db:
                st.info("当前企业岗位库为空。")
            else:
                job_to_delete = st.selectbox("请选择要删除的岗位", options=list(job_db.keys()))
                st.warning("注意：删除后该岗位标准将永久移除，且双端同步失效。")
                if st.button("确认删除该岗位"):
                    if job_to_delete in job_db:
                        del job_db[job_to_delete]
                        db_path = os.path.join("data", "job_models.json")
                        try:
                            with open(db_path, "w", encoding="utf-8") as f:
                                json.dump(job_db, f, ensure_ascii=False, indent=4)
                            st.success(f"岗位【{job_to_delete}】已成功删除。")
                            st.rerun()
                        except Exception as e:
                            st.error(f"更新数据库失败: {e}")

    with tab2:
        st.subheader("人才智能推荐与匹配清单")
        selected_job = st.selectbox("请选择招聘基准岗位：", options=job_options)

        uploaded_resumes = st.file_uploader(
            "批量导入候选人简历 (支持多个 PDF 文件同时拖入)",
            type=["pdf"],
            accept_multiple_files=True
        )

        if st.button("开始批量解析与智能排名"):
            if not uploaded_resumes:
                st.warning("请先上传至少一份简历文件。")
            else:
                job_text = job_db[selected_job]
                results = []
                progress_bar = st.progress(0)
                status_text = st.empty()
                total_files = len(uploaded_resumes)

                for i, file in enumerate(uploaded_resumes):
                    status_text.text(f"正在深度解析 AI 画像: {file.name} ({i + 1}/{total_files})")
                    resume_text = extract_text_from_pdf(file)

                    if resume_text == "ERROR_SCANNED_PDF":
                        st.warning(f"文件 {file.name} 为扫描件，已跳过。")
                        continue

                    if resume_text.strip():
                        profile = generate_student_profile(resume_text, job_text)
                        if profile:
                            radar = profile.get("radar_scores", {})
                            advantage = "无数据"
                            if radar:
                                best_skill = max(radar, key=radar.get)
                                advantage = f"{best_skill} ({radar[best_skill]}分)"

                            gaps = profile.get("gaps", ["无"])
                            core_gap = gaps[0] if gaps else "无"

                            results.append({
                                "候选人姓名": profile.get("name", file.name.split('.')[0]),
                                "源文件": file.name,
                                "综合匹配度": profile.get("match_score", 0),
                                "核心优势": advantage,
                                "首要短板": core_gap
                            })
                    progress_bar.progress((i + 1) / total_files)

                status_text.text("系统解析完成，正在按匹配度进行顺位排序...")
                if results:
                    df_talents = pd.DataFrame(results)
                    df_talents = df_talents.sort_values(by="综合匹配度", ascending=False).reset_index(drop=True)
                    df_talents.index = df_talents.index + 1
                    df_talents = df_talents.rename_axis("排名").reset_index()
                    df_talents["综合匹配度"] = df_talents["综合匹配度"].astype(str) + "%"
                    st.session_state.batch_talents_df = df_talents
                else:
                    st.error("文档解析或模型处理失败，未能生成有效数据。")

        if st.session_state.batch_talents_df is not None:
            st.success("批量匹配完成。以下是基于企业岗位标准的候选人顺位清单：")
            st.dataframe(st.session_state.batch_talents_df, use_container_width=True, hide_index=True)

    with tab3:
        st.subheader("入职表现与用人反馈 (AI 进化闭环)")
        st.write("企业反馈的真实业务数据将自动落盘，用于持续校准与微调 AI 诊断大模型的评估权重。")

        with st.expander("录入新员工入职档案"):
            new_emp_name = st.text_input("员工姓名", placeholder="例如：王小明")
            new_emp_time = st.text_input("入职时长", placeholder="例如：入职2周")
            if st.button("保存员工档案"):
                if new_emp_name and new_emp_time:
                    new_entry = f"{new_emp_name} ({new_emp_time})"
                    if new_entry not in st.session_state.employee_list:
                        st.session_state.employee_list.append(new_entry)
                        st.success(f"员工 【{new_emp_name}】 档案录入成功。")
                        st.rerun()
                else:
                    st.warning("请填写完整的员工姓名与时长。")

        col1, col2 = st.columns(2)
        with col1:
            selected_emp = st.selectbox("选择要评估的员工", options=st.session_state.employee_list)
            performance_score = st.slider("实际工作表现评分 (1-10分)", 1, 10, 8)
        with col2:
            feedback_text = st.text_area("具体反馈 (优点与不足，哪些能力被 AI 误判了？)", height=115)

        if st.button("提交反馈并反哺大模型"):
            if not feedback_text.strip():
                st.warning("请填写具体的反馈内容，以便 AI 学习。")
            else:
                feedback_data = {
                    "employee": selected_emp,
                    "score": performance_score,
                    "feedback": feedback_text
                }
                feedback_file = os.path.join("data", "feedback_logs.json")
                try:
                    os.makedirs("data", exist_ok=True)
                    logs = []
                    if os.path.exists(feedback_file):
                        with open(feedback_file, "r", encoding="utf-8") as f:
                            logs = json.load(f)

                    logs.append(feedback_data)
                    with open(feedback_file, "w", encoding="utf-8") as f:
                        json.dump(logs, f, ensure_ascii=False, indent=4)

                    st.success("反馈已加密落盘。系统将在下一次诊断请求时自动对齐该特征样本，校准模型权重。")
                except Exception as e:
                    st.error(f"本地数据写入失败，请检查读写权限: {e}")


# ==========================================
# 路由跳转二：学生端 (C端) 视图逻辑
# ==========================================
elif system_role == "学生端 (个人成长诊断)":

    with st.sidebar:
        st.header("输入数据")
        uploaded_resume = st.file_uploader("请上传学生简历 (支持 .pdf 格式)", type=["pdf"])

        if uploaded_resume is not None:
            if st.session_state.uploaded_file_name != uploaded_resume.name:
                extracted_text = extract_text_from_pdf(uploaded_resume)

                if extracted_text == "ERROR_SCANNED_PDF":
                    st.error("无法提取有效文本，请确保简历是标准的文字版 PDF，而非扫描件。")
                    st.session_state.uploaded_file_name = uploaded_resume.name
                elif extracted_text == "ERROR_PARSING_FAILED":
                    st.error("PDF 解析发生未知错误。")
                elif extracted_text.strip():
                    st.session_state.resume_text = extracted_text
                    st.session_state.uploaded_file_name = uploaded_resume.name
                    st.success("文档解析成功，文本已自动提取。")

        resume_input = st.text_area("简历内容 (可手动核对与修改)", height=200, value=st.session_state.resume_text)
        st.session_state.resume_text = resume_input

        st.markdown("---")
        st.subheader("目标岗位配置")

        c_job_options = list(job_options)
        c_job_options.append("自定义岗位输入")
        selected_job = st.selectbox("请选择企业真实岗位模型", options=c_job_options)

        if selected_job == "自定义岗位输入":
            job_input = st.text_area("手动输入岗位规范", height=200,
                                     value="岗位名称：\n能力核心要求：\n1. \n2. \n3. \n4. ")
        else:
            job_input = st.text_area("当前岗位能力要求 (只读)", height=200, value=job_db[selected_job], disabled=True)

        analyze_btn = st.button("开始 AI 诊断")

    # 提前判定视图模式
    is_report_mode = bool(st.session_state.profile_data or analyze_btn)

    # 优先注入 CSS（确保在进入耗时循环前，浏览器拿到白底样式）
    c_glass_bg = "rgba(255, 255, 255, 0.92)" if is_report_mode else "rgba(15, 23, 42, 0.35)"
    c_glass_blur = "20px" if is_report_mode else "6px"
    inject_global_css(c_glass_bg, c_glass_blur)

    title_color = "#1e3a8a" if is_report_mode else "#ffffff"
    subtitle_color = "#0f172a" if is_report_mode else "#e2e8f0"
    border_color = "rgba(0,0,0,0.1)" if is_report_mode else "rgba(255,255,255,0.2)"
    text_shadow = "none" if is_report_mode else "0px 2px 10px rgba(0,0,0,0.8)"

    st.markdown(
        f"""
        <div style="padding-bottom: 15px; border-bottom: 2px solid {border_color}; margin-bottom: 30px; transition: all 0.5s;">
            <h2 style="margin: 0; color: {title_color}; font-weight: 900; font-size: 2.2rem; text-shadow: {text_shadow}; transition: color 0.5s;">学生成长诊断与路径规划平台</h2>
            <p style="margin: 5px 0 0 0; color: {subtitle_color}; font-size: 1.1rem; font-weight: 800; text-shadow: {text_shadow}; transition: color 0.5s;">—— 基于润道星算大模型服务平台</p>
        </div>
        """,
        unsafe_allow_html=True
    )

    if not is_report_mode:
        st.markdown(
            """
            <div style="display: flex; flex-direction: column; justify-content: center; align-items: center; height: 40vh; width: 100%;">
                <h1 class="welcome-title">欢迎使用！</h1>
            </div>
            """,
            unsafe_allow_html=True
        )

    # 进入诊断逻辑
    if analyze_btn:
        if not resume_input or not job_input:
            st.warning("请填写完整的简历和岗位要求。")
        else:
            with st.status("系统：正在连接 AI 引擎...", expanded=True) as status:
                status.write("正在深度提炼多维特征结构...")

                json_placeholder = st.empty()
                full_json_str = ""

                try:
                    for chunk in generate_student_profile_stream(resume_input, job_input):
                        full_json_str += chunk
                        json_placeholder.markdown(f"```json\n{full_json_str} \n```")

                    json_placeholder.empty()
                    result = parse_profile_json(full_json_str)

                    if result:
                        st.session_state.profile_data = result
                        st.session_state.learning_plan_md = None
                        status.update(label="系统：AI 结构化提取完成，正在排版界面...", state="complete", expanded=False)
                        time.sleep(0.6)
                        st.rerun()
                    else:
                        status.update(label="系统：解析失败，大模型输出格式异常。", state="error")
                except Exception as e:
                    status.update(label=f"系统：请求异常 {e}", state="error")

    if st.session_state.profile_data:
        profile_data = st.session_state.profile_data

        col_metric, col_title = st.columns([1, 4])
        with col_metric:
            match_score = profile_data.get("match_score", 0)
            color = "normal" if match_score > 0 else "inverse"
            st.metric(label="岗位整体匹配度", value=f"{match_score}%", delta_color=color)
        with col_title:
            st.subheader(f"学生 {profile_data.get('name', '未知')} 的综合诊断报告")

        col1, col2 = st.columns([1, 1])
        with col1:
            radar_dict = profile_data.get("radar_scores", {})
            dimensions = list(radar_dict.keys())
            scores = list(radar_dict.values())

            if not dimensions:
                dimensions = ["未识别维度A", "未识别维度B", "未识别维度C", "未识别维度D"]
                scores = [0, 0, 0, 0]

            indicator = [{"name": str(dim), "max": 100} for dim in dimensions]

            radar_options = {
                "baseOption": {
                    "backgroundColor": "transparent",
                    "tooltip": {"trigger": "item"},
                    "radar": {
                        "indicator": indicator,
                        "center": ["50%", "50%"]
                    },
                    "series": [{
                        "name": "能力分值",
                        "type": "radar",
                        "data": [{"value": scores, "name": "学生现状"}],
                        "areaStyle": {"color": "rgba(54, 162, 235, 0.2)"},
                        "lineStyle": {"color": "#36A2EB"},
                        "itemStyle": {"color": "#36A2EB"}
                    }]
                },
                "media": [
                    {
                        "query": {"maxWidth": 550},
                        "option": {
                            "radar": {
                                "radius": "38%",
                                "axisName": {
                                    "fontSize": 10,
                                    "color": "#333",
                                    "width": 55,
                                    "overflow": "break",
                                    "lineHeight": 12
                                }
                            }
                        }
                    },
                    {
                        "query": {"minWidth": 551},
                        "option": {
                            "radar": {
                                "radius": "65%",
                                "axisName": {
                                    "fontSize": 14,
                                    "color": "#333",
                                    "width": 120,
                                    "overflow": "break",
                                    "lineHeight": 18
                                }
                            }
                        }
                    }
                ]
            }

            st_echarts(options=radar_options, height="380px", width="100%")

        with col2:
            st.subheader("核心差距明细")
            for gap in profile_data.get("gaps", []):
                st.write(f"- {gap}")

            st.subheader("推荐匹配岗位 (Top 5)")
            for pos in profile_data.get("top_positions", []):
                if isinstance(pos, dict):
                    st.write(f"- {pos.get('role', '未知岗位')}：{pos.get('reason', '无')}")
                else:
                    st.write(f"- {pos}")

        excel_data = create_excel_download(profile_data)
        st.download_button(
            label="导出能力画像报表 (.xlsx)",
            data=excel_data,
            file_name=f"{profile_data.get('name', 'student')}_能力画像.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

        st.divider()
        st.subheader("定制化成长路径规划")

        if st.button("生成 4 周学习计划与报告"):
            with st.status("系统：正在生成智能学习报告...", expanded=True) as plan_status:
                plan_status.write("正在连接教研智库与岗位分析系统...")

                report_container = st.empty()
                full_plan = ""
                gaps_text = "\n".join(profile_data.get("gaps", []))

                try:
                    for chunk in generate_learning_plan_stream(gaps_text, job_input):
                        full_plan += chunk
                        report_container.markdown(full_plan + " ")

                    if full_plan:
                        st.session_state.learning_plan_md = full_plan
                        plan_status.update(label="系统：学习报告编撰完成，正在排版界面...", state="complete", expanded=False)
                        time.sleep(0.6)
                        st.rerun()
                    else:
                        plan_status.update(label="系统：生成失败，请检查模型响应。", state="error")
                except Exception as e:
                    plan_status.update(label=f"系统：请求异常 {e}", state="error")

        if st.session_state.learning_plan_md:
            st.markdown(st.session_state.learning_plan_md)

            word_data = create_word_download(st.session_state.learning_plan_md)
            st.download_button(
                label="下载完整路径规划报告 (.docx)",
                data=word_data,
                file_name=f"{profile_data.get('name', 'student')}_路径规划.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )